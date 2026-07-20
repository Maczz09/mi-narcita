#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera el Informe de Resiliencia en Word (python-docx).
Fuente de evidencia: docs/informe-resiliencia/{shots,raw}. Estilo inspirado en
el Manual de Usuario de referencia (portada, secciones, figuras anotadas)."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHOTS = "docs/informe-resiliencia/shots"
RAW = "docs/informe-resiliencia/raw"
OUTP = "docs/informe-resiliencia/Informe_Resiliencia_NachoPps.docx"

NAVY = RGBColor(0x1E, 0x29, 0x3B)
RED = RGBColor(0xDC, 0x26, 0x26)
GREEN = RGBColor(0x15, 0x80, 0x3D)
AMBER = RGBColor(0xB4, 0x53, 0x09)
GREY = RGBColor(0x55, 0x5F, 0x6B)

doc = Document()

# ---------- estilos base ----------
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10.5)
for i in range(1, 4):
    h = doc.styles[f"Heading {i}"]
    h.font.color.rgb = NAVY
    h.font.name = "Calibri"

sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = Inches(0.9)
sec.top_margin = sec.bottom_margin = Inches(0.9)
CONTENT_W = 6.7


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def para(text="", size=10.5, bold=False, color=None, align=None, italic=False, space_after=6, style=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color:
            r.font.color.rgb = color
    return p


def runs(p, parts):
    """parts: lista de (texto, dict opts)"""
    for t, o in parts:
        r = p.add_run(t)
        r.font.size = Pt(o.get("size", 10.5))
        r.bold = o.get("bold", False); r.italic = o.get("italic", False)
        if o.get("mono"):
            r.font.name = "Consolas"
        if o.get("color"):
            r.font.color.rgb = o["color"]
    return p


def code_block(lines, cap=None):
    if cap:
        para(cap, size=9, italic=True, color=GREY, space_after=2)
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    c = t.rows[0].cells[0]
    shade(c, "0F172A")
    c.width = Inches(CONTENT_W)
    first = True
    for ln in lines:
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def figure(fname, caption, width=CONTENT_W):
    path = os.path.join(SHOTS, fname)
    if not os.path.exists(path):
        path = os.path.join(RAW, fname)
    if not os.path.exists(path):
        para(f"[falta figura {fname}]", color=RED); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width))
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(10)
    r = cp.add_run(caption); r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GREY


def table(headers, rows, widths, header_fill="1E293B"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i]); shade(c, header_fill)
        pp = c.paragraphs[0]; pp.paragraph_format.space_after = Pt(2)
        r = pp.add_run(h); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Inches(widths[i])
            pp = cells[i].paragraphs[0]; pp.paragraph_format.space_after = Pt(2)
            # val puede ser (texto,color)
            if isinstance(val, tuple):
                txt, col = val
            else:
                txt, col = val, None
            r = pp.add_run(txt); r.font.size = Pt(8.5)
            if col:
                r.font.color.rgb = col; r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def hrule():
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1E293B")
    pb.append(bottom); pPr.append(pb)


def pagebreak():
    doc.add_page_break()


# ================= PORTADA =================
for _ in range(3):
    doc.add_paragraph()
para("INFORME DE RESILIENCIA", size=13, bold=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Comportamiento del sistema ante la caída de cada servicio", size=24, bold=True,
     color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("NachoPps — Sistema Operativo de Restobar", size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("Pruebas de caos en vivo · 9 microservicios + API Gateway + PWA", size=11, bold=True,
     color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Qué falla · qué sigue funcionando · errores exactos · reintentos y tiempos · evidencia real",
     size=10, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
para("Fecha: 20 de julio de 2026", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Entorno: Docker Compose · Base URL http://localhost:8000 (Kong) · PWA http://localhost:4200",
     size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Método: apagado manual (docker stop) de cada contenedor, con evidencia en vivo "
     "(capturas PWA, red HTTP, logs estructurados y métricas Prometheus).",
     size=9, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
pagebreak()

# ================= CÓMO LEER =================
doc.add_heading("Cómo leer este informe", level=1)
para("Este documento responde, para cada servicio del sistema, a una pregunta operativa concreta: "
     "«si este servicio se cae ahora mismo, ¿qué pasa?». Cada sección se construyó apagando el "
     "contenedor real y observando el sistema en vivo.", space_after=6)
runs(doc.add_paragraph(), [
    ("Fuentes de evidencia usadas en cada sección: ", {"bold": True}),
])
for t, d in [
    ("Capturas de la PWA", " con anotaciones (círculos y recuadros rojos) sobre el punto exacto del fallo."),
    ("Red HTTP real", " — el código de estado que devolvió cada petición (200, 502, 503, 504…), capturado con Playwright."),
    ("Logs estructurados", " del servicio consumidor (líneas OperableLog: breaker, timeout, estado degradado)."),
    ("Métricas Prometheus", " — circuit_breaker_state, dependency_timeout_total, contadores de rechazo, /health/dependencies."),
]:
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(10)
    r2 = p.add_run(d); r2.font.size = Pt(10)
para("")
runs(para("", space_after=3), [("Convención de colores en las tablas: ", {"bold": True}),
     ("verde", {"color": GREEN, "bold": True}), (" = sigue funcionando · ", {}),
     ("rojo", {"color": RED, "bold": True}), (" = deja de funcionar · ", {}),
     ("ámbar", {"color": AMBER, "bold": True}), (" = degradado (funciona a medias, sin pérdida de datos).", {})])
hrule()

# ================= RESUMEN EJECUTIVO =================
doc.add_heading("Resumen ejecutivo", level=1)
para("El sistema está diseñado para degradar de forma acotada, no para caer en cascada. Ningún "
     "servicio caído tumba a los demás: cada dependencia síncrona está envuelta en timeout + circuit "
     "breaker + bulkhead, y los flujos de negocio se desacoplan por eventos (RabbitMQ + outbox). "
     "La ruta de dinero nunca pierde ni cobra dos veces.", space_after=8)

table(
    ["Servicio caído", "Radio de impacto", "Sigue funcionando", "Recuperación"],
    [
        ["Identidad", ("No se puede iniciar sesión ni refrescar token; Usuarios falla", RED),
         ("Sesiones ya abiertas operan hasta expirar el token (Kong valida JWT local)", GREEN), "Automática al volver"],
        ["Mesas", ("Pantalla Mesas no carga; pedido con mesa nueva → 503", RED),
         ("Pedido a mesa ya conocida (proyección local); resto de módulos", GREEN), "Breaker cierra ~30 s"],
        ["Pedidos", ("Pantallas Pedidos y Cocina no cargan", RED),
         ("Mesas, Caja, Inventario, Reservas, Reportes", GREEN), "Automática al volver"],
        ["Cuentas", ("Cobro rechazado (503) o cierre diferido", AMBER),
         ("Toda la pantalla Caja; el dinero queda consistente", GREEN), "Reconciliación + breaker"],
        ["Inventario", ("Pantallas Inventario/Compras vacías; pedido con producto nuevo → 503", RED),
         ("Pedido con producto ya conocido; resto de módulos", GREEN), "Breaker cierra ~30 s"],
        ["Caja", ("Pantalla Caja sin turno; no se cobra", RED),
         ("Mesas, Pedidos, Cocina, Inventario, Reservas, Reportes", GREEN), "Automática al volver"],
        ["Reservas", ("Pantalla Reservas vacía", RED), ("Todo el resto de la operación", GREEN), "Automática al volver"],
        ["Reportes", ("Pantalla Reportes e indicadores del dashboard vacíos", RED),
         ("Toda la operación real (mesas, pedidos, caja…)", GREEN), "Automática al volver"],
        ["Notificaciones", ("La campana no trae avisos nuevos", AMBER),
         ("TODA la operación de negocio (best-effort, no bloquea)", GREEN), "Outbox reenvía al volver"],
        ["Kong (gateway)", ("Nada opera: toda llamada a API falla", RED),
         ("La PWA se sigue sirviendo (nginx); datos ya cargados quedan visibles", AMBER), "Automática al volver"],
    ],
    [1.2, 2.3, 2.2, 1.0],
)
pagebreak()

# ================= MÉTODO Y ENTORNO =================
doc.add_heading("1. Método y entorno de prueba", level=1)
para("Todas las pruebas se ejecutaron contra el stack completo levantado con Docker Compose "
     "(9 microservicios NestJS, cada uno con su propia base PostgreSQL, más RabbitMQ, el API "
     "Gateway Kong, la PWA servida por nginx y la pila de observabilidad Prometheus/Grafana/Loki/"
     "Jaeger). El apagado de cada servicio se hizo con "
     , space_after=2)
runs(doc.add_paragraph(space_after=8) if False else doc.paragraphs[-1], [("docker stop <contenedor>", {"mono": True, "bold": True}),
     (", y el reinicio esperando el estado ", {}), ("healthy", {"mono": True}), (" antes de continuar con el siguiente.", {})])

doc.add_heading("1.1 Los nueve servicios y sus puertos", level=2)
table(
    ["Servicio", "Contenedor", "Puerto", "Responsabilidad", "Módulos PWA que dependen"],
    [
        ["identidad", "nachopps-servicio-identidad", "3001", "Login, JWT, usuarios", "Login, Usuarios"],
        ["mesas", "nachopps-servicio-mesas", "3002", "Estado del salón", "Mesas, Inicio"],
        ["pedidos", "nachopps-servicio-pedidos", "3004", "Comandas, KDS", "Pedidos, Cocina, Inicio"],
        ["cuentas", "nachopps-servicio-cuentas", "3005", "Cuentas por mesa, cierre", "Caja (cobro)"],
        ["reservas", "nachopps-servicio-reservas", "3006", "Reservas del día", "Reservas, Inicio"],
        ["inventario", "nachopps-servicio-inventario", "3007", "Productos y stock", "Inventario, Compras, Carta"],
        ["notificaciones", "nachopps-servicio-notificaciones", "3008", "Avisos, email async", "Campana"],
        ["caja", "nachopps-servicio-caja", "3009", "Turnos, cobros, cierre Z", "Caja, Inicio"],
        ["reportes", "nachopps-servicio-reportes", "3010", "KPIs, ventas", "Reportes, Inicio"],
    ],
    [1.0, 2.1, 0.55, 1.5, 1.55],
)
para("El API Gateway Kong (puerto 8000) es el único punto de entrada: valida el JWT localmente "
     "con la clave pública (no consulta a identidad en cada request), aplica rate-limiting y enruta "
     "a cada servicio. La PWA (puerto 4200) es una SPA servida por nginx, independiente de Kong.",
     space_after=6)

doc.add_heading("1.2 Cómo el sistema absorbe una caída (mapa mental)", level=2)
for t, d in [
    ("Llamada directa de la PWA a un servicio caído", " → Kong no alcanza el upstream y responde 502/504. La PWA lo traduce a un banner y reintenta."),
    ("Llamada entre servicios", " (p. ej. pedidos→inventario) → el cliente HTTP tiene timeout + circuit breaker + bulkhead; al fallar devuelve 503 controlado, nunca cuelga."),
    ("Flujo de negocio asíncrono", " (stock, cuenta automática, notificaciones) → viaja por RabbitMQ con outbox; si el consumidor está caído, el evento espera en la cola y se procesa al volver."),
]:
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(10)
    p.add_run(d).font.size = Pt(10)
pagebreak()

# ================= MARCO DE RESILIENCIA =================
doc.add_heading("2. Marco de resiliencia (los números exactos)", level=1)
para("Esta sección fija los valores que se repiten en todo el informe: cuántos reintentos, cada "
     "cuánto tiempo, con qué timeout y cuándo abre el circuito. El código es la fuente de verdad "
     "(libs/resiliencia, libs/observabilidad y los clientes HTTP); aquí está el resumen operativo.", space_after=8)

doc.add_heading("2.1 En el navegador (PWA)", level=2)
table(
    ["Aspecto", "Valor", "Detalle"],
    [
        ["Timeout por petición", "8 s", "AbortSignal.timeout por intento"],
        ["Reintentos de red (cliente)", "2", "solo GET o POST con Idempotency-Key; ante 5xx / error de red"],
        ["Backoff", "300·2ⁿ ms + jitter", "tope 2 s entre intentos"],
        ["Reintentos de React Query", "2 (salvo 404)", "sobre cada query de lectura"],
        ["Re-consulta en error", "cada 5 s", "refetchInterval mientras la query siga en error"],
        ["Caché offline", "24 h", "las lecturas GET se persisten en localStorage (último dato conocido)"],
        ["401 → refresh", "1 intento", "si el refresh falla → cierre de sesión"],
    ],
    [1.7, 1.5, 3.5],
)
para("Traducción de errores que ve el usuario:", bold=True, space_after=2)
table(
    ["Situación", "Mensaje en pantalla"],
    [
        [("502 / 503 / 504 (servicio caído tras Kong)", None), ("«El servicio no está disponible en este momento.»", RED)],
        [("Sin respuesta / timeout de red (Kong caído)", None), ("«El servidor no responde. Verifica tu conexión o inténtalo de nuevo.»", RED)],
        [("429 (rate-limit)", None), ("«Demasiadas solicitudes. Espera un momento.»", AMBER)],
    ],
    [3.3, 3.4],
)

doc.add_heading("2.2 Entre servicios (backend)", level=2)
table(
    ["Dependencia (llamada)", "Timeout", "Reintentos", "Breaker", "Al fallar"],
    [
        ["pedidos → mesas (leer mesa)", "2 s", "1 (GET)", "2,5 s", ("503 «El servicio de mesas no responde»", RED)],
        ["pedidos → inventario (validar stock)", "2 s", ("0 (POST lote)", RED), "2,5 s", ("503 «El servicio de inventario no responde»", RED)],
        ["caja → cuentas (leer cuenta)", "2 s", "1 (GET)", "2,5 s", ("503 «No se pudo obtener la cuenta»", RED)],
        ["caja → cuentas (cerrar = dinero)", "4 s", ("0", RED), "4,5 s", ("degrada a PAGO_SIN_CIERRE_CONFIRMADO", AMBER)],
    ],
    [2.1, 0.75, 1.05, 0.7, 2.1],
)
runs(para("", space_after=2), [
    ("Circuit breaker: ", {"bold": True}),
    ("umbral de error 50 %, resetTimeout 30 s, los 4xx no abren el circuito. Estado observable en ",
     {}), ("circuit_breaker_state", {"mono": True}),
    (" → 0 = cerrado, 0.5 = medio-abierto, 1 = abierto.", {})])
runs(para("", space_after=2), [
    ("Bulkhead: ", {"bold": True}),
    ("pool de 10 llamadas concurrentes por dependencia; al saturarse hace shed-load (rechaza) en vez de encolar sin límite.", {})])
runs(para("", space_after=8), [
    ("RabbitMQ (eventos): ", {"bold": True}),
    ("reintento con backoff exponencial + jitter (3 intentos: 250·2ⁿ⁻¹ ms) y luego DLQ. "
     "Un mensaje envenenado (payload corrupto) va directo a la DLQ al primer intento.", {})])

runs(para("", space_after=8, size=9.5, italic=True), [
    ("Detalle clave que se repite abajo: ", {"bold": True, "italic": True, "color": NAVY}),
    ("crear un pedido NO llama a mesas/inventario si la mesa y los productos ya están en la "
     "proyección local del servicio de pedidos. La llamada remota (y por tanto el breaker) solo se "
     "dispara en «cold-start»: una mesa o un producto que pedidos todavía no conoce. Por eso, con "
     "mesas o inventario caídos, se puede seguir tomando pedidos de lo ya conocido y solo fallan los "
     "ítems nuevos.", {"italic": True})])
pagebreak()

# ================= MATRIZ =================
doc.add_heading("3. Matriz servicio caído × módulo de la PWA", level=1)
para("Lectura rápida de todo el informe en una tabla. Cada fila es un servicio apagado; cada celda, "
     "el efecto sobre ese módulo.", space_after=8)
OK = ("OK", GREEN); DOWN = ("cae", RED); DEG = ("degr.", AMBER); NA = ("—", GREY)
table(
    ["Apago ↓ / Módulo →", "Login", "Inicio", "Mesas", "Pedidos", "Cocina", "Caja", "Invent.", "Reservas", "Reportes"],
    [
        ["Identidad", DOWN, DEG, DEG, DEG, DEG, DEG, DEG, DEG, DEG],
        ["Mesas", OK, DEG, DOWN, DEG, OK, OK, OK, OK, OK],
        ["Pedidos", OK, DEG, OK, DOWN, DOWN, OK, OK, OK, OK],
        ["Cuentas", OK, OK, OK, OK, OK, DEG, OK, OK, OK],
        ["Inventario", OK, DEG, OK, DEG, OK, OK, DOWN, OK, OK],
        ["Caja", OK, DEG, OK, OK, OK, DOWN, OK, OK, OK],
        ["Reservas", OK, DEG, OK, OK, OK, OK, OK, DOWN, OK],
        ["Reportes", OK, DEG, OK, OK, OK, OK, OK, OK, DOWN],
        ["Notificac.", OK, OK, OK, OK, OK, OK, OK, OK, OK],
        ["Kong", DOWN, DOWN, DOWN, DOWN, DOWN, DOWN, DOWN, DOWN, DOWN],
    ],
    [1.15, 0.62, 0.62, 0.62, 0.66, 0.62, 0.58, 0.66, 0.72, 0.72],
)
runs(para("", size=9, space_after=2), [("OK", {"color": GREEN, "bold": True}), (" funciona · ", {}),
     ("cae", {"color": RED, "bold": True}), (" no funciona · ", {}),
     ("degr.", {"color": AMBER, "bold": True}),
     (" degradado: la pantalla abre pero muestra ceros/vacío o una acción puntual falla.", {})])
para("«degr.» en el dashboard de Inicio significa que la tarjeta que alimenta ese servicio queda en "
     "0 o vacía, sin cartel de error (ver Hallazgo transversal 1).", size=9, italic=True, color=GREY)
pagebreak()

# ================= SECCIONES POR SERVICIO =================
doc.add_heading("4. Detalle por servicio", level=1)
para("Cada servicio sigue la misma plantilla: qué depende de él, cómo se apagó, qué deja de "
     "funcionar y qué sigue, los errores exactos, qué ocurre en el background (reintentos y tiempos) "
     "y la evidencia en vivo.", space_after=8)


def svc_header(num, titulo, contenedor, puerto):
    doc.add_heading(f"4.{num}  {titulo}", level=2)
    runs(para("", space_after=6), [
        ("Contenedor: ", {"bold": True}), (contenedor, {"mono": True}),
        ("   ·   Puerto: ", {"bold": True}), (puerto, {"mono": True}),
        ("   ·   Apagado: ", {"bold": True}), (f"docker stop {contenedor}", {"mono": True})])


# ---- 4.1 MESAS ----
svc_header(1, "Mesas", "nachopps-servicio-mesas", "3002")
para("Gobierna el estado del salón. La pantalla Mesas lo lee directo; el dashboard de Inicio usa su "
     "conteo de ocupación; y crear un pedido necesita resolver la mesa (solo si es una mesa que "
     "pedidos aún no conoce).", space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Pantalla Mesas", None), ("Cae", RED), "Spinner perpetuo «Cargando salón…» con esqueletos; reintenta cada 5 s", ("502", RED)],
       [("Inicio · tarjeta Mesas", None), ("Degrada", AMBER), "«Mesas ocupadas 0/0 · 0 % de ocupación» (sin cartel de error)", ("—", GREY)],
       [("Crear pedido a mesa ya conocida", None), ("Funciona", GREEN), "Pedido creado normal (usa proyección local)", ("201", GREEN)],
       [("Crear pedido a mesa NUEVA", None), ("Cae", RED), "«El servicio de mesas no responde. Reintente.»", ("503", RED)],
      ], [1.75, 0.75, 3.15, 0.5])
runs(para("", space_after=6), [
    ("En el background: ", {"bold": True}),
    ("la pantalla Mesas nunca llegó a tener datos, así que se queda en «Cargando» reintentando "
     "indefinidamente cada 5 s (se observaron 24 respuestas 502 en 26 s). El breaker de pedidos→mesas "
     "tiene timeout 2 s + 1 reintento (peor caso ≈ 4,25 s) antes de devolver 503; al superar el "
     "umbral abre el circuito 30 s.", {})])
figure("mesas_02_modulo.png", "Fig. 4.1a — Pantalla Mesas con el servicio caído: «Cargando salón…» "
       "perpetuo (1). Nota: el indicador «En línea» sigue verde (2) porque mide la conexión al gateway, no la salud del servicio.")
figure("mesas_01_inicio.png", "Fig. 4.1b — Dashboard de Inicio: la tarjeta «Mesas ocupadas» degrada a 0/0 "
       "en silencio (1), sin banner de error. Comparar con el baseline (64/144).")
code_block([
    'POST /v1/pedidos (mesa cold, mesas caído)  →  HTTP 503 en 2 558 ms',
    '{"errorCode":"HTTP_503","message":"El servicio de mesas no responde. Reintente."}',
    '',
    '# /health/dependencies de pedidos:  "mesas":"DOWN"',
    '# Prometheus:  circuit_breaker_open_total{mesas}=2  dependency_timeout_total{mesas}=2  retry_attempts_total=2',
], cap="Evidencia en vivo (breaker pedidos→mesas):")
hrule()

# ---- 4.2 PEDIDOS ----
svc_header(2, "Pedidos", "nachopps-servicio-pedidos", "3004")
para("Es el corazón operativo: comandas y KDS (cocina). Las pantallas Pedidos y Cocina lo leen "
     "directo. No hay otro servicio que lo llame de forma síncrona, así que su caída no arrastra a "
     "mesas, caja ni inventario.", space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Pantalla Pedidos", None), ("Cae", RED), "«Cargando…» perpetuo; reintenta cada 5 s", ("502", RED)],
       [("Pantalla Cocina (KDS)", None), ("Cae", RED), "«Cargando…» perpetuo", ("502", RED)],
       [("Inicio · widgets de pedidos", None), ("Degrada", AMBER), "«Sin actividad reciente · Todo en orden. Sin alertas.»", ("—", GREY)],
       [("Mesas / Caja / Inventario / Reservas", None), ("Funciona", GREEN), "Operación normal", ("200", GREEN)],
      ], [1.75, 0.75, 3.15, 0.5])
runs(para("", space_after=6), [
    ("Riesgo operativo a vigilar: ", {"bold": True, "color": AMBER}),
    ("en el dashboard, la caída de pedidos se ve como «Todo en orden, sin alertas» — es decir, "
     "parece calma cuando en realidad el sistema está ciego a los pedidos. Es el caso más engañoso "
     "de degradación silenciosa.", {})])
figure("pedidos_02.png", "Fig. 4.2 — Pantalla Pedidos con el servicio caído: «Cargando…» perpetuo (1), reintentando cada 5 s.")
hrule()

# ---- 4.3 CUENTAS ----
svc_header(3, "Cuentas (ruta de dinero)", "nachopps-servicio-cuentas", "3005")
para("Lleva la cuenta abierta por mesa y su cierre. La caja lo llama en el cobro: primero lee la "
     "cuenta y luego la cierra. Es la ruta crítica de dinero, y por eso su diseño es el más "
     "conservador del sistema: nunca reintenta el cierre y degrada de forma honesta.", space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Pantalla Caja (ver turno/movimientos)", None), ("Funciona", GREEN), "La pantalla abre y muestra todo (lee del servicio caja, que sigue vivo)", ("200", GREEN)],
       [("Cobrar cuenta (cuentas totalmente caído)", None), ("Cae", RED), "«No se pudo obtener la cuenta. Reintente.» — el pago NO se registra", ("503", RED)],
       [("Cobrar cuenta (cuentas cae a mitad del cobro)", None), ("Degrada", AMBER), "«Pago registrado; cierre de cuenta en proceso» — el pago SÍ queda", ("201", GREEN)],
      ], [2.15, 0.8, 3.0, 0.5])
runs(para("", space_after=4), [
    ("Por qué el dinero está a salvo en los dos casos: ", {"bold": True}),
    ("(a) si cuentas está caído del todo, el cobro se rechaza en la lectura, antes de tocar dinero. "
     "(b) Si cuentas cae después de leer pero antes de cerrar, el pago ya se persistió (transacción "
     "local + outbox) y la cuenta queda en estado ", {}),
    ("PAGO_SIN_CIERRE_CONFIRMADO", {"mono": True}),
    ("; un servicio de reconciliación la cierra sola cuando cuentas vuelve. Nunca se cobra dos veces "
     "ni se pierde el pago.", {})])
figure("cuentas_02_caja.png", "Fig. 4.3 — La pantalla Caja funciona con cuentas caído (2, franja «Caja abierta»); "
       "solo falla la acción «Cobrar cuenta» (1), que devuelve 503 sin registrar el pago.")
code_block([
    'POST /v1/caja/pagos (cuentas caído)  →  HTTP 503 en 2 573 ms',
    '{"errorCode":"HTTP_503","message":"No se pudo obtener la cuenta. Reintente."}',
    '',
    '# log de caja (OperableLog):',
    'CircuitBreaker  "Llamada vencida por timeout del circuito."  op=CuentasHttpClient.fetchCuentaConBreaker',
    'CircuitBreaker  circuitBreakerState:"OPEN"  "Circuito abierto; rechazando peticiones a la dependencia."',
    '# caso mid-flow captado en vivo:',
    'CierreReconciliacionService  resultingState:"PAGO_SIN_CIERRE_CONFIRMADO"  errorCode:"CIERRE_RECONCILIACION_PENDIENTE"',
], cap="Evidencia en vivo (ruta de dinero):")
hrule()

# ---- 4.4 INVENTARIO ----
svc_header(4, "Inventario", "nachopps-servicio-inventario", "3007")
para("Productos, precios y stock. Las pantallas Inventario, Compras y Carta lo leen. Además, crear "
     "un pedido valida el stock contra inventario, pero solo para productos que pedidos aún no "
     "conoce (cold-start).", space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Pantalla Inventario", None), ("Cae", RED), "Tarjetas en 0 (Disponibles/Stock bajo/Agotados) y tabla en esqueleto", ("502", RED)],
       [("Pantalla Compras", None), ("Cae", RED), "No carga el catálogo", ("502", RED)],
       [("Crear pedido con producto ya conocido", None), ("Funciona", GREEN), "Pedido creado (usa proyección local)", ("201", GREEN)],
       [("Crear pedido con producto NUEVO", None), ("Cae", RED), "«El servicio de inventario no responde. Reintente.»", ("503", RED)],
      ], [2.0, 0.8, 3.15, 0.5])
runs(para("", space_after=6), [
    ("En el background: ", {"bold": True}),
    ("la validación de stock es un POST de lote y por diseño NO se reintenta (no es idempotente sin "
     "clave), así que falla rápido: timeout 2 s → 503 y suma a ", {}),
    ("pedidos_rechazados_dependencia_total{inventario}", {"mono": True}), (".", {})])
figure("inventario_02.png", "Fig. 4.4 — Pantalla Inventario con el servicio caído: contadores en 0 (1) "
       "y «0 productos» (2) con la tabla en esqueleto. El formulario «Nuevo producto» sigue visible pero fallaría al guardar.")
code_block([
    'POST /v1/pedidos (producto cold, inventario caído)  →  HTTP 503 en 2 098 ms',
    '{"errorCode":"HTTP_503","message":"El servicio de inventario no responde. Reintente."}',
    '',
    '# log de pedidos (OperableLog):',
    'AppService          resultingState:"1 productos en cold-start"  "cargando desde inventario"',
    'CircuitBreaker      circuitBreakerState:"OPEN"  op=InventarioHttpClient.fetchProductosLote',
    'InventarioHttpClient durationMs:2013  errorCode:"DEPENDENCY_TIMEOUT"  "pedido rechazado por dependencia caída"',
], cap="Evidencia en vivo (breaker pedidos→inventario):")
hrule()

# ---- 4.5 CAJA ----
svc_header(5, "Caja", "nachopps-servicio-caja", "3009")
para("Turnos, cobros y cierre Z. La pantalla Caja y la tarjeta de caja del dashboard lo leen. Su "
     "caída no arrastra a otros servicios (es consumidor de eventos, no dependencia síncrona de nadie).",
     space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Pantalla Caja", None), ("Cae", RED), "«Sin turno abierto»; ventas, movimientos y métodos en 0", ("504", RED)],
       [("Inicio · tarjeta de caja", None), ("Degrada", AMBER), "«caja sin turno»", ("—", GREY)],
       [("Cobrar / abrir-cerrar turno", None), ("Cae", RED), "No se puede operar la caja", ("504", RED)],
       [("Resto de módulos", None), ("Funciona", GREEN), "Operación normal", ("200", GREEN)],
      ], [1.9, 0.8, 3.2, 0.55])
figure("caja_02.png", "Fig. 4.5 — Pantalla Caja con el servicio caído: degrada a «Sin turno abierto» (1), "
       "con todos los importes y contadores en cero.")
hrule()

# ---- 4.6 RESERVAS ----
svc_header(6, "Reservas", "nachopps-servicio-reservas", "3006")
para("Reservas del día. Solo la pantalla Reservas y un indicador del dashboard dependen de él; su "
     "caída es la de menor radio de impacto sobre la operación de salón.", space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Pantalla Reservas", None), ("Cae", RED), "«Sin reservas para hoy» (estado vacío — puede confundirse con «no hay reservas»)", ("502", RED)],
       [("Inicio · indicador de reservas", None), ("Degrada", AMBER), "vacío", ("—", GREY)],
       [("Resto de módulos", None), ("Funciona", GREEN), "Operación normal", ("200", GREEN)],
      ], [1.9, 0.8, 3.2, 0.55])
figure("reservas_02.png", "Fig. 4.6 — Pantalla Reservas con el servicio caído (1): degrada a estado vacío «Sin reservas para hoy».")
hrule()

# ---- 4.7 REPORTES ----
svc_header(7, "Reportes", "nachopps-servicio-reportes", "3010")
para("KPIs y ventas del día. Alimenta la pantalla Reportes y la tarjeta «Ventas del día» del "
     "dashboard. No participa en ningún flujo de negocio, así que su caída no afecta a mesas, "
     "pedidos ni caja reales.", space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Pantalla Reportes", None), ("Cae", RED), "Tarjetas en esqueleto (cargando); reintenta cada 5 s", ("502", RED)],
       [("Inicio · tarjeta Ventas del día", None), ("Degrada", AMBER), "«Sin ventas registradas hoy»", ("—", GREY)],
       [("Operación real (mesas/pedidos/caja)", None), ("Funciona", GREEN), "Intacta", ("200", GREEN)],
      ], [2.0, 0.8, 3.1, 0.55])
figure("reportes_02.png", "Fig. 4.7 — Pantalla Reportes con el servicio caído (1): tarjetas en esqueleto de carga.")
hrule()

# ---- 4.8 NOTIFICACIONES ----
svc_header(8, "Notificaciones", "nachopps-servicio-notificaciones", "3008")
para("Avisos en la campana y envío de emails de forma asíncrona. Es best-effort por diseño: está "
     "desacoplado por RabbitMQ + outbox, así que su caída NO bloquea ningún flujo de negocio.",
     space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Campana de notificaciones", None), ("Degrada", AMBER), "No trae avisos nuevos (los ya cargados siguen)", ("504", RED)],
       [("Inicio y TODA la operación", None), ("Funciona", GREEN), "Mesas, pedidos, cobros, inventario, reservas: 100 % normal", ("200", GREEN)],
       [("Emails / avisos generados durante la caída", None), ("Diferido", AMBER), "Quedan en el outbox y se envían al volver el servicio", ("—", GREY)],
      ], [2.15, 0.8, 3.0, 0.5])
runs(para("", space_after=8), [
    ("Conclusión: ", {"bold": True, "color": GREEN}),
    ("es el único servicio cuya caída no tiene impacto operativo. El dashboard de Inicio se "
     "renderizó completo y correcto con notificaciones apagado; solo la campana devolvió 504.", {})])
hrule()

# ---- 4.9 IDENTIDAD ----
svc_header(9, "Identidad", "nachopps-servicio-identidad", "3001")
para("Login, emisión y refresco de JWT, y gestión de usuarios. Es especial porque Kong valida el "
     "token localmente con la clave pública: no consulta a identidad en cada petición. Por eso una "
     "sesión ya abierta puede seguir leyendo de otros servicios aunque identidad esté caído.",
     space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Iniciar sesión (nuevo)", None), ("Cae", RED), "«El servicio no está disponible en este momento.»", ("504", RED)],
       [("Recargar la pestaña estando logueado", None), ("Cae", RED), "La app valida con /auth/me → falla → vuelve al login", ("502", RED)],
       [("Sesión ya abierta (sin recargar)", None), ("Funciona", GREEN), "Sigue operando con el token en memoria hasta que expire (~15 min)", ("200", GREEN)],
       [("Pantalla Usuarios (crear/editar)", None), ("Cae", RED), "No se pueden gestionar usuarios", ("502/504", RED)],
      ], [2.2, 0.8, 2.95, 0.55])
runs(para("", space_after=6), [
    ("Matiz importante: ", {"bold": True}),
    ("el token de acceso vive en memoria de la pestaña. Mientras la pestaña siga abierta, las "
     "lecturas y escrituras a otros servicios funcionan (Kong valida el JWT sin llamar a identidad). "
     "Pero cualquier recarga o expiración del token necesita ", {}),
    ("/auth/refresh", {"mono": True}),
    (" o ", {}), ("/auth/me", {"mono": True}),
    (", que sí dependen de identidad; al fallar, la sesión cae al login.", {})])
figure("identidad_01_login.png", "Fig. 4.9a — Login con identidad caído: banner «El servicio no está disponible "
       "en este momento» (1). El intento de login reintenta (refresh 504, login 504 ×3).")
figure("identidad_02_sesion_viva.png", "Fig. 4.9b — Recargar la pestaña con identidad caído: la validación de "
       "sesión (/auth/me) falla con 502 y la app vuelve al login (1); no llega a cargar ningún módulo.")
hrule()

# ---- 4.10 KONG ----
svc_header(10, "Kong (API Gateway) — caída total", "nachopps-kong", "8000")
para("Kong es el único punto de entrada a las APIs. Si cae, ningún servicio es alcanzable, aunque "
     "todos sigan vivos por dentro. La PWA, en cambio, la sirve nginx (puerto 4200), así que la "
     "aplicación se sigue abriendo — pero sin datos.", space_after=6)
table(["Módulo / petición", "Qué pasa", "Qué ve el usuario", "HTTP"],
      [
       [("Cargar la PWA", None), ("Funciona", GREEN), "La app abre (la sirve nginx, no Kong)", ("200", GREEN)],
       [("Iniciar sesión", None), ("Cae", RED), "«El servidor no responde. Verifica tu conexión o inténtalo de nuevo.»", ("red", RED)],
       [("Cualquier módulo con datos", None), ("Cae", RED), "Toda petición a API falla por conexión rechazada", ("red", RED)],
       [("Datos ya cargados antes de la caída", None), ("Degrada", AMBER), "La caché offline (24 h) muestra el último dato conocido", ("—", GREY)],
      ], [2.1, 0.8, 3.05, 0.55])
runs(para("", space_after=6), [
    ("Diferencia clave con «identidad caído»: ", {"bold": True}),
    ("aquí las peticiones ni siquiera reciben respuesta HTTP (conexión rechazada), por lo que el "
     "banner es distinto — «El servidor no responde» en vez de «El servicio no está disponible». "
     "Es la señal de que el problema es el gateway/red, no un servicio puntual.", {})])
figure("kong_01_login.png", "Fig. 4.10 — Login con Kong caído: banner de fallo de red «El servidor no responde. "
       "Verifica tu conexión o inténtalo de nuevo» (1). La PWA se sirvió igual (nginx), pero no hay backend alcanzable.")
pagebreak()

# ================= HALLAZGOS TRANSVERSALES =================
doc.add_heading("5. Hallazgos transversales", level=1)
findings = [
    ("Degradación silenciosa en los dashboards",
     "Ante un fallo de lectura, las tarjetas de Inicio (y pantallas como Reportes o Reservas) "
     "muestran 0 o «sin datos» en vez de un cartel de error. Es robusto (la app no se rompe) pero "
     "puede confundir «servicio caído» con «no hay actividad». El caso más delicado es pedidos: el "
     "dashboard dice «Todo en orden, sin alertas» mientras el sistema está ciego a los pedidos."),
    ("El indicador «En línea» no refleja la salud por servicio",
     "El punto verde «En línea» del encabezado mide la conectividad al gateway, no si cada "
     "microservicio responde. Puede seguir verde con un servicio caído."),
    ("Spinner perpetuo vs. estado vacío (inconsistencia de UX)",
     "Algunas pantallas sin dato previo se quedan en «Cargando…» reintentando cada 5 s (Mesas, "
     "Pedidos, Cocina, Reportes/Inventario en esqueleto); otras degradan a un estado vacío limpio "
     "(Reservas «Sin reservas»). Convendría unificar hacia un estado de error explícito y accionable."),
    ("Dos banners de error distintos, y está bien que así sea",
     "«El servicio no está disponible» (HTTP 5xx desde Kong) señala un servicio puntual caído; «El "
     "servidor no responde» (sin respuesta de red) señala Kong/red. La distinción ayuda a "
     "diagnosticar sin abrir la consola."),
    ("La ruta de dinero nunca pierde ni duplica",
     "Cobro con cuentas caído: si falla la lectura, se rechaza limpio (503, nada se cobra); si falla "
     "el cierre, el pago queda registrado una sola vez y la cuenta se reconcilia sola al volver "
     "(PAGO_SIN_CIERRE_CONFIRMADO). El cierre nunca se reintenta automáticamente para no arriesgar doble cobro."),
    ("Las proyecciones locales dan resiliencia real",
     "Crear pedidos sigue funcionando con mesas o inventario caídos siempre que la mesa y los "
     "productos ya se conozcan (proyección local sincronizada por eventos). Solo el «cold-start» "
     "—algo nuevo que el servicio de pedidos no vio— dispara la llamada remota y, con la dependencia "
     "caída, el 503."),
    ("Caché offline de 24 h",
     "Las lecturas GET se persisten en el navegador; una pestaña que ya cargó datos los sigue "
     "mostrando (marcados como último dato conocido) aunque el backend o Kong caigan después."),
]
for i, (t, d) in enumerate(findings, 1):
    runs(para("", space_after=2), [(f"{i}. {t}", {"bold": True, "size": 11, "color": NAVY})])
    para(d, space_after=8)
hrule()

# ================= APÉNDICE =================
doc.add_heading("6. Apéndice — cómo reproducir", level=1)
para("Todas las capturas y evidencias de este informe se generaron con scripts versionados en "
     "scripts/resiliencia-shots/. El patrón por servicio es:", space_after=6)
code_block([
    '# 1) apagar el servicio y ver el impacto en la PWA + red',
    'docker stop nachopps-servicio-<svc>',
    'node scripts/resiliencia-shots/shoot.mjs <out> <Modulo> - 12000   # captura + red',
    '',
    '# 2) snapshot de observabilidad (breaker + /health/dependencies)',
    'bash scripts/resiliencia-shots/snap.sh <svc>_down',
    '',
    '# 3) ejercitar un flujo entre servicios con la sesión real (cookie) de la PWA',
    'node scripts/resiliencia-shots/flow.mjs \'[{"method":"POST","path":"/v1/pedidos","body":{...}}]\'',
    '',
    '# 4) restaurar y esperar healthy',
    'docker start nachopps-servicio-<svc>',
], cap="Flujo de captura por servicio:")
para("Referencias de implementación (fuente de verdad): libs/resiliencia (breaker, bulkhead, retry, "
     "retry de broker), libs/observabilidad (métricas, prefetch), los clientes HTTP "
     "(servicio-pedidos→inventario/mesas, servicio-caja→cuentas), y la política en "
     "docs/resiliencia-politica.md y docs/matriz-resiliencia.md.", size=9.5, italic=True, color=GREY)

doc.save(OUTP)
print("OK ->", OUTP)

